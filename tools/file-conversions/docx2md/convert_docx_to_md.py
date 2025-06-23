from docx import Document
import markdownify

def convert_docx_to_markdown(docx_path, md_path):
    # Load the Word document
    doc = Document(docx_path)
    
    # Initialize an empty list to store lines of Markdown content
    md_lines = []
    
    # Function to convert a single paragraph to Markdown
    def paragraph_to_markdown(para):
        if para.style.name.startswith('Heading'):
            level = int(para.style.name[7:])  # Assuming Heading 1, Heading 2, etc.
            return '\n' + f"{'#' * level} {markdownify.markdownify(para.text)}" + '\n'
        elif para.style.name == 'List Paragraph':
            # Handle bullet points
            return  '\n' + f"- {markdownify.markdownify(para.text)}" + '\n'
        else:
            return '\n' + markdownify.markdownify(para.text) + '\n'
    
    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        md_text = paragraph_to_markdown(para)
        md_lines.append(md_text)
    
    # Function to convert a table to Markdown
    def table_to_markdown(table):
        md_table = []
        header_row = '| ' + ' | '.join(cell.text for cell in table.rows[0].cells) + ' |'
        separator = '|' + '|'.join(['---'] * len(table.columns)) + '|'
        md_table.append(header_row)
        md_table.append(separator)
        
        for row in table.rows[1:]:
            md_row = '| ' + ' | '.join(cell.text for cell in row.cells) + ' |'
            md_table.append(md_row)
        
        return '\n' + '\n'.join(md_table) + '\n'
    
    # Iterate through each table in the document
    for table in doc.tables:
        md_lines.append(table_to_markdown(table))
    
    # Write the Markdown content to a file
    with open(md_path, 'w', encoding='utf-8') as md_file:
        md_file.write('\n'.join(md_lines))

# Example usage
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Convert Word document to Markdown.")
    parser.add_argument("docx_path", help="Path to the input Word document (.docx)")
    parser.add_argument("md_path", help="Path to the output Markdown file (.md)")
    args = parser.parse_args()

    convert_docx_to_markdown(args.docx_path, args.md_path)
    print(f"Converted {args.docx_path} to {args.md_path}")
